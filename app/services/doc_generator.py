from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from typing import Dict, Any
import uuid


class DocxGenerator:
    def __init__(self, template: Dict[str, Any], data: Dict[str, Any]):
        """
        Инициализация генератора документов
        
        :param template: JSON-шаблон документа
        :param data: Данные для заполнения плейсхолдеров
        """
        self.template = template
        self.data = data
        self.document = Document()
        self._setup_page_settings()
    
    def _setup_page_settings(self):
        """Настройка параметров страницы из шаблона"""
        if 'page_settings' not in self.template:
            return
            
        settings = self.template['page_settings']
        sections = self.document.sections
        for section in sections:
            # Ориентация страницы
            if settings.get('orientation') == 'landscape':
                section.orientation = 1  # Landscape
            
            # Поля страницы
            margins = settings.get('margins', {})
            section.top_margin = Inches(margins.get('top', 1) / 2.54)
            section.bottom_margin = Inches(margins.get('bottom', 1) / 2.54)
            section.left_margin = Inches(margins.get('left', 1) / 2.54)
            section.right_margin = Inches(margins.get('right', 1) / 2.54)
            
            # Шрифт по умолчанию
            font = settings.get('font', {})
            if 'name' in font or 'size' in font:
                style = self.document.styles['Normal']
                font_obj = style.font
                font_obj.name = font.get('name', 'Times New Roman')
                font_obj.size = Pt(font.get('size', 12))
    
    def _replace_placeholders(self, text: str) -> str:
        """Замена плейсхолдеров в тексте на реальные значения"""
        if not isinstance(text, str):
            return text
            
        for key, value in self.data.items():
            placeholder = f"{{{key}}}"
            text = text.replace(placeholder, str(value))
        return text
    
    def _apply_paragraph_style(self, paragraph, style: Dict[str, Any]):
        """Применение стилей к параграфу"""
        if not style:
            return
            
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        
        # Жирность/курсив
        if style.get('bold'):
            run.bold = True
        if style.get('italic'):
            run.italic = True
            
        # Выравнивание
        align = style.get('alignment')
        if align == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        # Размер шрифта
        if style.get('font_size'):
            run.font.size = Pt(style['font_size'])
            
        # Отступ после
        if style.get('spacing_after'):
            paragraph.paragraph_format.space_after = Pt(style['spacing_after'])
    
    def _add_document_number(self):
        """Добавление номера документа"""
        if 'document_number' not in self.template:
            return
            
        config = self.template['document_number']
        method = config.get('generation_method', 'auto_increment')
        
        if method == 'auto_increment':
            doc_num = f"{config.get('prefix', '')}{str(uuid.uuid4())[:8].upper()}"
        else:
            doc_num = config.get('prefix', '')
            
        paragraph = self.document.add_paragraph(f"№ {doc_num}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.document.add_paragraph()







        
    
    def _process_header(self, element: Dict[str, Any]):
        """Обработка заголовка документа"""
        text = self._replace_placeholders(element.get('text', ''))
        paragraph = self.document.add_paragraph(text)
        self._apply_paragraph_style(paragraph, element.get('style', {}))
    
    def _process_paragraph(self, element: Dict[str, Any]):
        """Обработка обычного параграфа"""
        text = self._replace_placeholders(element.get('text', ''))
        paragraph = self.document.add_paragraph(text)
        self._apply_paragraph_style(paragraph, element.get('style', {}))
    
    def _process_section(self, element: Dict[str, Any]):
        """Обработка раздела документа"""
        # Добавляем заголовок раздела
        title = self._replace_placeholders(element.get('title', ''))
        paragraph = self.document.add_paragraph(title)
        
        # Применяем стили заголовка
        style = element.get('style', {})
        if style.get('bold'):
            paragraph.runs[0].bold = True
            
        # Обрабатываем содержимое раздела
        for content_item in element.get('content', []):
            self._process_content_element(content_item)
    
    def _process_clause(self, element: Dict[str, Any]):
        """Обработка нумерованного пункта с ручной нумерацией"""
        number = element.get('number', '')
        text = self._replace_placeholders(element.get('text', ''))
        
        paragraph = self.document.add_paragraph()
        paragraph.add_run(f"{number}. ").bold = True  # Номер жирный
        paragraph.add_run(text)
        
        self._apply_paragraph_style(paragraph, element.get('style', {}))
    
    def _process_list(self, element: Dict[str, Any]):
        """Обработка списка"""
        for item in element.get('items', []):
            text = self._replace_placeholders(item)
            paragraph = self.document.add_paragraph(text, style='ListBullet')
            self._apply_paragraph_style(paragraph, element.get('style', {}))
    
    def _process_table(self, element: Dict[str, Any]):
        """Обработка таблицы"""
        # Создаем таблицу
        rows = element.get('rows', [])
        cols = element.get('columns', [])
        table = self.document.add_table(rows=len(rows) + 1, cols=len(cols))
        
        # Настройки выравнивания таблицы
        align = element.get('alignment')
        if align == 'center':
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        elif align == 'right':
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        
        # Добавляем заголовки
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(cols):
            hdr_cells[i].text = col
            if element.get('bold_headers'):
                hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Добавляем строки данных
        for i, row in enumerate(rows, start=1):
            for j, cell in enumerate(row):
                table.rows[i].cells[j].text = self._replace_placeholders(cell)
                if align == 'center':
                    table.rows[i].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _process_signature_line(self, element: Dict[str, Any]):
        """Обработка строки для подписи"""
        text = self._replace_placeholders(element.get('text', ''))
        paragraph = self.document.add_paragraph()
        paragraph.add_run(text)
        self._apply_paragraph_style(paragraph, element.get('style', {}))








    
    
    def _process_content_element(self, element: Dict[str, Any]):
        """Обработка элемента контента по его типу"""
        element_type = element.get('type')
        
        processors = {
            'header': self._process_header,
            'paragraph': self._process_paragraph,
            'section': self._process_section,
            'clause': self._process_clause,
            'list': self._process_list,
            'table': self._process_table,
            'signature_line': self._process_signature_line
        }
        
        if element_type in processors:
            processors[element_type](element)
        else:
            print(f"Warning: Unknown element type '{element_type}'")
    
    def generate(self) -> 'DocxGenerator':
        """Генерация документа на основе шаблона"""
        # Добавляем номер документа
        self._add_document_number()
        
        # Обрабатываем основной контент
        for element in self.template.get('content', []):
            self._process_content_element(element)
        
        return self
    
    def save_to_file(self, filename: str) -> None:
        """Сохранение документа в файл"""
        self.generate()
        self.document.save(filename)
    
    def save_to_bytesio(self) -> BytesIO:
        """Сохранение документа в поток байтов (для FastAPI)"""
        self.generate()
        file_stream = BytesIO()
        self.document.save(file_stream)
        file_stream.seek(0)
        return file_stream
    
    def get_document(self) -> Document:
        """Получение объекта Document для дополнительных манипуляций"""
        self.generate()
        return self.document